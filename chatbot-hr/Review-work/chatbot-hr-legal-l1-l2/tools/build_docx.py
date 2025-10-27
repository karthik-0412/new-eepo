from docx import Document
from pathlib import Path

md = Path(__file__).parents[1] / 'DOCLING.md'
assert md.exists(), f"{md} not found"

text = md.read_text(encoding='utf-8')

# naive markdown to docx: split by headers and paragraphs
doc = Document()
lines = text.splitlines()
for line in lines:
    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
    elif line.strip().startswith('```'):
        # code block start/end
        # collect all lines until closing ```
        # simple approach: skip the ``` line and add a paragraph with monospaced font
        pass
    else:
        doc.add_paragraph(line)

out = Path(__file__).parents[1] / 'DOCLING.docx'
doc.save(out)
print('Wrote', out)
